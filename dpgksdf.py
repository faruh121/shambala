from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

# Создаем документ Word
doc = Document()

# Добавляем заголовок задания
doc.add_paragraph("Задание 9. Создайте таблицу как представлено ниже")

# Создаем таблицу с 6 столбцами и 4 строками
table = doc.add_table(rows=4, cols=6)
table.style = 'Table Grid'

# Настраиваем ширину столбцов
table.columns[0].width = Pt(120)
table.columns[1].width = Pt(80)
table.columns[2].width = Pt(100)
table.columns[3].width = Pt(100)
table.columns[4].width = Pt(100)
table.columns[5].width = Pt(100)

# Заполняем заголовки
row = table.rows[0]
row.cells[0].merge(row.cells[1])  # Объединяем ячейки
row.cells[0].text = "Имя файла и его расширение"
row.cells[2].text = "Исходный размер"
cell_3 = row.cells[3].merge(row.cells[5])  # Объединяем три ячейки
row.cells[3].text = "Размер после архивирования (степень сжатия)"

# Вставляем строки данных и заголовки под ними
row2 = table.rows[1]
row2.cells[2].text = "WinRar"
row2.cells[3].text = "Zip"

doc.save("Table_A9.docx")